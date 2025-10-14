import streamlit as st
import pandas as pd
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# Set page config
st.set_page_config(
    page_title="Quotation Generator",
    page_icon="🏭",
    layout="wide"
)

# Custom CSS for styling
st.markdown("""
<style>
    .main-header {
        text-align: center;
        font-size: 5em;
        font-weight: bold;
        color: #F2F0EF;
        margin-bottom: 10px;
    }
    .doc-header {
        text-align: center;
        font-size: 2.5em;
        font-weight: bold;
        color: #e74c3c;
        text-decoration: italic;
        margin: 20px 0;
    }
    .total-amount {
        font-size: 1.2em;
        font-weight: bold;
        color: #000000;
        background-color: #f8f9fa;
        padding: 10px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .section-header {
        font-size: 1.1em;
        font-weight: bold;
        color: #b9c1bb;
        margin-top: 20px;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'items' not in st.session_state:
    st.session_state["items"] = []

# Main header
st.markdown('<div class="doc-header">Quotation Generator</div>', unsafe_allow_html=True)

# Sidebar for main inputs
with st.sidebar:
    st.header("Basic Information")
    
    # Business name input (customizable)
    business_name = st.text_input("Business/Organization Name", value="MASIH TRADERS")

    #GSTIN
    gstin = st.text_input("Enter GSTIN",value = "09AMIPM2416L1ZT")

    #Contact
    contact = st.text_input("Enter Contact Number", value = "9839 710 370")

    # Customer name
    customer_name = st.text_input("Customer Name", placeholder="Enter customer name")
    
    # Document type
    doc_type = st.selectbox("Document Type", ["QUOTATION"])
    
    # Date
    doc_date = st.date_input("Date", datetime.now())
    
    # Format date as DD/MM/YYYY
    formatted_date = doc_date.strftime("%d/%m/%Y")

# Main content area
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown('<div class="section-header">Items</div>', unsafe_allow_html=True)
    
    # Item input form
    with st.expander("➕ Add New Item", expanded=True):
        item_col1, item_col2 = st.columns(2)
        
        with item_col1:
            item_name = st.text_input("Item Name", placeholder="e.g., MS Gate")
            item_spec = st.text_input("Specification", placeholder="e.g., 96\" Height X 63\" Breadth")
            
        with item_col2:
            # Quantity with unit
            qty_col1, qty_col2 = st.columns([2, 1])
            with qty_col1:
                quantity = st.number_input("Quantity", min_value=0.0, step=0.1)
            with qty_col2:
                unit = st.selectbox("Unit", ["kg", "gm", "pcs", "box", "packet", "sqft", "ft", "inch", "quintal", "ton", "custom"])
                if unit == "custom":
                    custom_unit = st.text_input("Custom Unit", placeholder="Enter unit")
                    unit = custom_unit if custom_unit else "pcs"
            
            rate = st.number_input("Rate per Unit", min_value=0.0, step=0.01)
        
        # Calculate total for this item
        item_total = quantity * rate
        st.write(f"**Item Total: ₹{item_total:,.2f}**")
        
        # Add item button
        if st.button("Add Item", type="primary"):
            if item_name and quantity > 0 and rate > 0:
                new_item = {
                    "name": item_name,
                    "specification": item_spec,
                    "quantity": quantity,
                    "unit": unit,
                    "rate": rate,
                    "total": item_total
                }
                st.session_state["items"].append(new_item)
                st.success(f"Added: {item_name}")
                st.rerun()
            else:
                st.error("Please fill all required fields")

    # Display added items
    if st.session_state["items"]:
        st.markdown('<div class="section-header">Added Items</div>', unsafe_allow_html=True)
        
        for i, item in enumerate(st.session_state["items"]):
            col_item, col_remove = st.columns([4, 1])
            
            with col_item:
                st.write(f"**{i+1}. {item['name']}:** {item['specification']} -- {item['quantity']} {item['unit']} @{item['rate']} = ₹{item['total']:,.2f}")
            
            with col_remove:
                if st.button("❌", key=f"remove_{i}"):
                    st.session_state["items"].pop(i)
                    st.rerun()
        
        # Calculate total amount
        total_amount = sum(item['total'] for item in st.session_state["items"])
        st.markdown(f'<div class="total-amount">Total Amount: ₹{total_amount:,.2f}</div>', unsafe_allow_html=True)

with col2:
    st.markdown('<div class="section-header">💰 Payment Details</div>', unsafe_allow_html=True)
    
    # Payment details
    advance = st.number_input("Advance Amount (₹)", min_value=0.0, step=0.01)
    
    # # Calculate balance but if advance exceeds total amount then balance becomes zero
    total_amount = sum(item['total'] for item in st.session_state["items"]) if st.session_state["items"] else 0
    balance = abs(total_amount - advance)
    # Compare advance with total amount
    if advance > total_amount:
        overpaid = advance - total_amount
        balance = overpaid
    else:
        overpaid = 0
        balance = total_amount - advance

    
    st.write(f"**Balance Amount: ₹{balance:,.2f}**")

    # Calculate balance
    # st.write(
    #     f"**Balance Amount: ₹{sum(item['total'] for item in st.session_state['items']) - advance:,.2f}**")
    
    # Material used
    st.markdown('<div class="section-header">🔧 Material Used</div>', unsafe_allow_html=True)
    material_type = st.selectbox("Material Type", ["MS (Mild Steel)", "SS (Stainless Steel)"])
    material_grade = st.text_area("Material Grade/Detail", placeholder="e.g., 16 g - 304")
    
    # Additional materials
    additional_materials = st.text_area("Section Used:", 
                                       placeholder="e.g., Apollo Pipe: 2x4\" (30 kg)\nBidding: 25x5\"\nCNC Sheet: 3MM",
                                       height=100)
    
    # Terms and conditions
    st.markdown('<div class="section-header">📜 Terms & Conditions</div>', unsafe_allow_html=True)
    
    terms_options = {
        "advance_payment": "75% Advance Payment",
        "workshop_payment": "25% Payment at workshop before delivery",
        "gst_applicable": "18% GST applicable on online payments, invoice will be provided",
        "pit_digging": "Pit digging before gate installation is customer's responsibility",
        "primer_not_included": "Primer application is not included",
        "fiber_sheet_extra": "Fiber Sheet charges will be extra",
        "cartage_extra": "Cartage charges will be extra",
        "black_gate_design": "Quotation based on Black Gate Design",
        "completion_time": "The work is said to be completed within 30 days after the date of the advance payment"
    }
    
    selected_terms = []
    for key, term in terms_options.items():
        if st.checkbox(term, key=key):
            selected_terms.append(term)
    
    # Custom terms
    custom_terms = st.text_area("Custom Terms", placeholder="Add any custom terms and conditions")
    if custom_terms:
        selected_terms.extend(custom_terms.split('\n'))

# Generate document section
st.markdown("---")
st.markdown('<div class="section-header">📄 Generate Document</div>', unsafe_allow_html=True)

col_preview, col_download = st.columns(2)

with col_preview:
    if st.button("📋 Preview Document", type="primary"):
        if customer_name and st.session_state["items"]:
            st.markdown("---")
            st.markdown(f"**{business_name}**")
            st.markdown(f"**[{doc_type}] Date: {formatted_date}**")
            st.markdown(f"**To,**  \n**{customer_name}**")
            
            # Items list
            for i, item in enumerate(st.session_state["items"], 1):
                spec_text = f": {item['specification']}" if item['specification'] else ""
                st.markdown(f"{i}. **{item['name']}{spec_text}** -- {item['quantity']} {item['unit']} @{item['rate']} {item['total']:.0f}/-")
            
            st.markdown(f"**Total Amount** {total_amount:.0f}/-")
            
            if advance > 0:
                st.markdown(f"**Advance** {advance:.0f}/-")
                st.markdown(f"**Balance** {balance:.0f}/-")
                
            
            # Material used
            if additional_materials:
                st.markdown("**Material Used:**")
                for material in additional_materials.split('\n'):
                    if material.strip():
                        st.markdown(f"- **{material.strip()}**")
            
            # Terms and conditions
            if selected_terms:
                st.markdown("**Terms & Conditions:**")
                for term in selected_terms:
                    if term.strip():
                        st.markdown(f"- **{term.strip()}**")
        else:
            st.error("Please enter customer name and add at least one item")

with col_download:
    def create_word_document():
        doc = Document()
        
        # Set page margins
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # Header with GSTIN and Contact
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(f'GSTIN: {gstin}		                                                                                    Contact: +91 {contact}')
        header_run.font.name = 'Calibri'
        header_run.font.size = Pt(12)
        header_run.bold = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Main heading
        main_heading = doc.add_paragraph()
        main_heading_run = main_heading.add_run(business_name)
        main_heading_run.font.name = 'Adobe Garamond Pro Bold'
        main_heading_run.font.size = Pt(40)
        main_heading_run.bold = True
        main_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # # 🔧 Adjust spacing below heading
        main_heading.paragraph_format.space_after = Pt(9) # reduce space below
        
        # # Underline (line of underscores)
        # underline_para = doc.add_paragraph()
        # underline_run = underline_para.add_run('_' * 98)
        # underline_run.font.name = 'Calibri'
        # underline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Document type and date in separate paragraphs
        doc_type_para = doc.add_paragraph()
        doc_type_run = doc_type_para.add_run(doc_type)
        doc_type_run.font.name = 'Calibri'
        doc_type_run.font.size = Pt(24)
        doc_type_run.underline = True
        doc_type_run.bold = True
        doc_type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'Date: {formatted_date}')
        date_run.font.name = 'Calibri Light (Headings)'
        date_run.font.size = Pt(16)
        date_run.bold = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Customer info
        customer_para = doc.add_paragraph()

        # "To,"
        customer_run1 = customer_para.add_run("To,")
        customer_run1.font.size = Pt(16)
        customer_run1.font.name = "Calibri"
        customer_run1.bold = True

        # Customer name
        customer_run2 = customer_para.add_run(f"\n{customer_name}")
        customer_run2.font.size = Pt(20)
        customer_run2.font.name = "Calibri"
        customer_run2.bold = True
        
        # Create table for items with columns
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Item'
        header_cells[1].text = 'Qty'
        header_cells[2].text = 'Rate'
        header_cells[3].text = 'Amount'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Calibri'
        
        # Add items to table
        for i, item in enumerate(st.session_state["items"], 1):
            row_cells = table.add_row().cells
            
            # Item name and specification (indented)
            spec_text = f": {item['specification']}" if item['specification'] else ""
            item_text = f"{i}. {item['name']}{spec_text}"
            row_cells[0].text = item_text
            
            # Quantity with unit
            row_cells[1].text = f"{item['quantity']} {item['unit']}"
            
            # Rate
            row_cells[2].text = f"{item['rate']}"
            
            # Amount
            row_cells[3].text = f"{item['total']:.0f}/-"
        
        # Add total, advance, and balance rows
        # Total Amount
        total_row = table.add_row().cells
        total_row[0].text = ""
        total_row[1].text = ""
        total_row[2].text = "Total Amount"
        total_row[3].text = f"{total_amount:.0f}/-"
        
        # Make total row bold
        for cell in total_row[2:]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Advance (if any)
        if advance > 0:
            advance_row = table.add_row().cells
            advance_row[0].text = ""
            advance_row[1].text = ""
            advance_row[2].text = "Advance"
            advance_row[3].text = f"{advance:.0f}/-"
            # Make advance row bold
            for cell in advance_row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Balance
            balance_row = table.add_row().cells
            balance_row[0].text = ""
            balance_row[1].text = ""
            balance_row[2].text = "Balance"
            balance_row[3].text = f"{balance:.0f}/-"
            # Make balance row bold
            for cell in balance_row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        # Material used
        if additional_materials:
            material_full = f"{material_type} {material_grade}" if material_grade else material_type
            material_header = doc.add_paragraph()
            material_header_run = material_header.add_run(f'Material Used: {material_full}')
            material_header_run.bold = True
            material_header_run.font.name = 'Calibri'
            material_header_run.font.size = Pt(14)
            
            for material in additional_materials.split('\n'):
                if material.strip():
                    material_para = doc.add_paragraph()
                    material_run = material_para.add_run(f"• {material.strip()}")
                    material_run.font.name = 'Calibri'
                    material_run.font.size = Pt(14)

        
        # Terms and conditions
        if selected_terms:
            terms_header = doc.add_paragraph()
            terms_header_run = terms_header.add_run('Terms & Conditions:')
            terms_header_run.bold = True
            terms_header_run.font.name = 'Calibri'
            terms_header_run.font.size = Pt(14)
            
            for term in selected_terms:
                if term.strip():
                    term_para = doc.add_paragraph()
                    term_run = term_para.add_run(f"• {term.strip()}")
                    term_run.font.name = 'Calibri'
                    term_run.font.size = Pt(14)
        
        return doc
    
    if st.button("📥 Download Word Document", type="secondary"):
        if customer_name and st.session_state["items"]:
            doc = create_word_document()
            
            # Save to bytes
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Download button
            st.download_button(
                label="📄 Download Document",
                data=doc_buffer.getvalue(),
                file_name=f"{doc_type}_{customer_name}_{formatted_date.replace('/', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Please enter customer name and add at least one item")

# Clear all button
if st.button("🗑️ Clear All Data", type="secondary"):
    st.session_state["items"] = []
    st.rerun()

# Footer
st.markdown("---")
st.markdown("*Quotation Generator v1.0*")