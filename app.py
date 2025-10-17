import streamlit as st
import pandas as pd
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# Set page config
st.set_page_config(
    page_title="Quotation Generator",
    page_icon="🏭",
    layout="wide"
)

# Custom CSS for styling
st.markdown("""
<style>
    .main-header {
        text-align: center;
        font-size: 5em;
        font-weight: bold;
        color: #F2F0EF;
        margin-bottom: 10px;
    }
    .doc-header {
        text-align: center;
        font-size: 2.5em;
        font-weight: bold;
        color: #e74c3c;
        text-decoration: italic;
        margin: 20px 0;
    }
    .total-amount {
        font-size: 1.2em;
        font-weight: bold;
        color: #000000;
        background-color: #f8f9fa;
        padding: 10px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .section-header {
        font-size: 1.1em;
        font-weight: bold;
        color: #b9c1bb;
        margin-top: 20px;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'items' not in st.session_state:
    st.session_state["items"] = []

# Main header
st.markdown('<div class="doc-header">Quotation Generator</div>', unsafe_allow_html=True)

# Sidebar for main inputs
with st.sidebar:
    st.header("Basic Information")
    
    # Business name input (customizable)
    business_name = st.text_input("Business/Organization Name", value="MASIH TRADERS")

    #GSTIN
    gstin = st.text_input("Enter GSTIN", value="09AMIPM2416L1ZT")

    #Contact
    contact = st.text_input("Enter Contact Number", value="9839 710 370")

    # Customer name
    customer_name = st.text_input("Customer Name", placeholder="Enter customer name")
    
    # Document type
    doc_type = st.selectbox("Document Type", ["QUOTATION"])
    
    # Date
    doc_date = st.date_input("Date", datetime.now())
    
    # Format date as DD/MM/YYYY
    formatted_date = doc_date.strftime("%d/%m/%Y")

# Main content area
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown('<div class="section-header">Items</div>', unsafe_allow_html=True)
    
    with st.expander("➕ Add New Item", expanded=True):
        item_name = st.text_input("Item Name", placeholder="e.g., MS Gate", key="item_name")
        item_spec = st.text_input("Specification", placeholder='e.g., 96" Height X 63" Breadth', key="item_spec")

        st.markdown("**Select Entry Mode:**")
        entry_mode = st.radio(
            "Select how you want to enter item details:",
            ["Quantity + Rate", "Quantity + Amount", "Rate + Amount", "Amount Only"],
            horizontal=True,
            key="entry_mode"
        )

        # Default values
        quantity = 0.0
        rate = 0.0
        amount = 0.0
        unit = ""

        if entry_mode in ["Quantity + Rate", "Quantity + Amount"]:
            qty_col1, qty_col2 = st.columns([2, 1])
            with qty_col1:
                quantity = st.number_input("Quantity", min_value=0.0, step=0.1, key="item_qty")
            with qty_col2:
                unit = st.selectbox(
                    "Unit",
                    ["kg", "gm", "pcs", "box", "packet", "sqft", "ft", "inch", "quintal", "ton", "custom"],
                    key="item_unit"
                )
                if unit == "custom":
                    custom_unit = st.text_input("Custom Unit", placeholder="Enter unit", key="item_custom_unit")
                    unit = custom_unit if custom_unit else "pcs"

        if entry_mode == "Quantity + Rate":
            rate = st.number_input("Rate per Unit (₹)", min_value=0.0, step=0.01, key="item_rate")
            amount = quantity * rate

        elif entry_mode == "Quantity + Amount":
            amount = st.number_input("Amount (₹)", min_value=0.0, step=0.01, key="item_amount_q")

        elif entry_mode == "Rate + Amount":
            rate = st.number_input("Rate (₹)", min_value=0.0, step=0.01, key="item_rate_ra")
            amount = st.number_input("Amount (₹)", min_value=0.0, step=0.01, key="item_amount_ra")

        elif entry_mode == "Amount Only":
            amount = st.number_input("Amount (₹)", min_value=0.0, step=0.01, key="item_amount_only")

        st.write(f"**Item Total: ₹{amount:,.2f}**")

        # Add item
        if st.button("Add Item", type="primary", key="add_item_btn"):
            new_item = {
                "name": item_name if item_name else "—",
                "specification": item_spec if item_spec else "",
                "quantity": quantity,
                "unit": unit,
                "rate": rate,
                "total": amount,
                "mode": entry_mode
            }
            st.session_state["items"].append(new_item)
            st.success(f"Added: {new_item['name']}")
            st.rerun()

    # Display added items
    if st.session_state["items"]:
        st.markdown('<div class="section-header">Added Items</div>', unsafe_allow_html=True)
        
        for i, item in enumerate(st.session_state["items"]):
            col_item, col_remove = st.columns([4, 1])
            
            with col_item:
                mode = item.get("mode", "Quantity + Rate")
                spec_text = f": {item['specification']}" if item.get('specification') else ""
                
                if mode == "Amount Only":
                    st.write(f"**{i+1}. {item['name']}{spec_text}** -- ₹{item['total']:,.2f}")
                elif mode == "Quantity + Amount":
                    st.write(f"**{i+1}. {item['name']}{spec_text}** -- {item['quantity']} {item['unit']} = ₹{item['total']:,.2f}")
                elif mode == "Rate + Amount":
                    st.write(f"**{i+1}. {item['name']}{spec_text}** -- @₹{item['rate']:.2f} = ₹{item['total']:,.2f}")
                else:  # Quantity + Rate
                    st.write(f"**{i+1}. {item['name']}{spec_text}** -- {item['quantity']} {item['unit']} @₹{item['rate']:.2f} = ₹{item['total']:,.2f}")
            
            with col_remove:
                if st.button("❌", key=f"remove_{i}"):
                    st.session_state["items"].pop(i)
                    st.rerun()
        
        # Calculate total amount
        total_amount = sum(item['total'] for item in st.session_state["items"])
        st.markdown(f'<div class="total-amount">Total Amount: ₹{total_amount:,.2f}</div>', unsafe_allow_html=True)

with col2:
    st.markdown('<div class="section-header">💰 Payment Details</div>', unsafe_allow_html=True)

    # Advance
    advance = st.number_input("Advance Amount (₹)", min_value=0.0, step=0.01, key="advance")

    # GST controls
    st.markdown('<div class="section-header">🧾 GST</div>', unsafe_allow_html=True)
    gst_choice = st.selectbox("GST Rate", ["None", "0%", "5%", "12%", "18%", "28%"], index=4, key="gst_choice")
    gst_mode = st.radio("GST Mode", ["Exclusive (to be added)", "Inclusive (amount includes GST)"], index=0, key="gst_mode")

    # Subtotal from items
    subtotal = sum(item['total'] for item in st.session_state["items"]) if st.session_state["items"] else 0.0
    st.write(f"**Subtotal:** ₹{subtotal:,.2f}")

    # compute GST values
    gst_percent = 0.0 if gst_choice == "None" else float(gst_choice.replace("%", ""))
    if gst_percent > 0:
        if gst_mode.startswith("Exclusive"):
            gst_amount = subtotal * gst_percent / 100.0
            taxable_base = subtotal
            amount_chargeable = subtotal + gst_amount
        else:  # Inclusive: amount already includes GST
            gst_amount = subtotal * gst_percent / (100.0 + gst_percent)
            taxable_base = subtotal - gst_amount
            amount_chargeable = subtotal
        cgst = gst_amount / 2.0
        sgst = gst_amount / 2.0

        st.write(f"**CGST @{gst_percent/2:.2f}%:** ₹{cgst:,.2f}")
        st.write(f"**SGST @{gst_percent/2:.2f}%:** ₹{sgst:,.2f}")
        st.markdown(f"**Amount Chargeable:** ₹{amount_chargeable:,.2f}")
    else:
        gst_amount = cgst = sgst = 0.0
        amount_chargeable = subtotal
        st.write(f"**Amount Chargeable:** ₹{amount_chargeable:,.2f}")

    # Balance after advance
    balance = max(0.0, amount_chargeable - advance)
    st.write(f"**Balance Amount: ₹{balance:,.2f}**")
    
    # -----------------------------
    # MATERIALS (multiple sections)
    # -----------------------------
    st.markdown('<div class="section-header">🔧 Material Used</div>', unsafe_allow_html=True)

    # Ensure materials list exists
    if "materials" not in st.session_state:
        st.session_state["materials"] = []  # each entry: {"type":..., "grade":..., "details":[...]}

    # Expander to add a new material section
    with st.expander("➕ Add Material Section", expanded=False):
        mat_type = st.selectbox("Material Type", ["MS (Mild Steel)", "SS (Stainless Steel)", "Aluminium"], key="mat_type_input")
        mat_grade = st.text_input("Material Grade/Detail", placeholder="e.g., 16 g - 304", key="mat_grade_input")
        mat_details = st.text_area(
            "Section Used (one item per line)",
            placeholder='e.g., 1" Sq Pipe\n0.5" Sq Pipe\nCNC Sheet: 3MM',
            height=120,
            key="mat_details_input"
        )

        if st.button("Add Material Section", key="add_material_btn"):
            details_lines = [line.strip() for line in mat_details.split("\n") if line.strip()]
            if not details_lines:
                st.error("Please enter at least one material detail (one per line).")
            else:
                st.session_state["materials"].append({
                    "type": mat_type,
                    "grade": mat_grade.strip(),
                    "details": details_lines
                })
                st.success(f"Added material section: {mat_type} {mat_grade}".strip())
                st.rerun()

    # Display added material sections (with remove option)
    if st.session_state["materials"]:
        st.markdown("### Added Material Sections")
        for idx, mat in enumerate(st.session_state["materials"]):
            with st.expander(f"{idx+1}. {mat['type']} {mat['grade']}".strip(), expanded=False):
                st.write("**Details:**")
                for d in mat["details"]:
                    st.write(f"- {d}")
                cols = st.columns([1, 1, 3])
                with cols[0]:
                    if st.button("Edit", key=f"edit_mat_{idx}"):
                        # Prefill inputs for quick edit (simple approach)
                        st.session_state["mat_edit_index"] = idx
                        st.session_state["mat_type_input"] = mat["type"]
                        st.session_state["mat_grade_input"] = mat["grade"]
                        st.session_state["mat_details_input"] = "\n".join(mat["details"])
                        st.rerun()
                with cols[1]:
                    if st.button("❌ Remove", key=f"remove_mat_{idx}"):
                        st.session_state["materials"].pop(idx)
                        st.success("Removed.")
                        st.rerun()
                with cols[2]:
                    st.markdown("<small>Use Edit to load this section back into the Add form for quick changes.</small>", unsafe_allow_html=True)


# Generate document section
st.markdown("---")
st.markdown('<div class="section-header">📄 Generate Document</div>', unsafe_allow_html=True)

col_preview, col_download = st.columns(2)

with col_preview:
    if st.button("📋 Preview Document", type="primary", key="preview"):
        st.markdown("---")
        st.markdown(f"**{business_name}**")
        st.markdown(f"**[{doc_type}] Date: {formatted_date}**")
        st.markdown(f"**To,**  \n**{customer_name or '—'}**")

        # Items list
        if st.session_state["items"]:
            for i, item in enumerate(st.session_state["items"], 1):
                mode = item.get("mode", "Quantity + Rate")
                spec_text = f": {item['specification']}" if item.get('specification') else ""
                
                if mode == "Amount Only":
                    st.markdown(f"{i}. **{item['name']}{spec_text}** = ₹{item['total']:,.2f}")
                elif mode == "Quantity + Amount":
                    st.markdown(f"{i}. **{item['name']}{spec_text}** -- {item['quantity']} {item['unit']} = ₹{item['total']:,.2f}")
                elif mode == "Rate + Amount":
                    st.markdown(f"{i}. **{item['name']}{spec_text}** -- @₹{item['rate']:,.2f} = ₹{item['total']:,.2f}")
                else:  # Quantity + Rate
                    st.markdown(f"{i}. **{item['name']}{spec_text}** -- {item['quantity']} {item['unit']} @₹{item['rate']:,.2f} = ₹{item['total']:,.2f}")
        else:
            st.markdown("_No items added._")

        # Subtotal
        subtotal = sum(item['total'] for item in st.session_state["items"]) if st.session_state["items"] else 0.0
        st.markdown(f"**Subtotal:** ₹{subtotal:,.2f}")

        # GST
        gst_percent = 0.0 if gst_choice == "None" else float(gst_choice.replace("%", ""))
        if gst_percent > 0:
            if gst_mode.startswith("Exclusive"):
                gst_amount = subtotal * gst_percent / 100
                taxable_base = subtotal
                amount_chargeable = subtotal + gst_amount
            else:  # Inclusive
                gst_amount = subtotal * gst_percent / (100 + gst_percent)
                taxable_base = subtotal - gst_amount
                amount_chargeable = subtotal

            cgst = gst_amount / 2
            sgst = gst_amount / 2

            st.markdown(f"**CGST @{gst_percent/2:.2f}%:** ₹{cgst:,.2f}")
            st.markdown(f"**SGST @{gst_percent/2:.2f}%:** ₹{sgst:,.2f}")
            st.markdown(f"**Amount Chargeable:** ₹{amount_chargeable:,.2f}")
        else:
            amount_chargeable = subtotal
            st.markdown(f"**Amount Chargeable:** ₹{amount_chargeable:,.2f}")

        # Advance & Balance
        if advance > 0:
            st.markdown(f"**Advance:** ₹{advance:,.2f}")
            st.markdown(f"**Balance:** ₹{max(0, amount_chargeable - advance):,.2f}")

with col_download:
    def create_word_document():
        doc = Document()
        
        # Set page margins
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # Header with GSTIN and Contact
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(f'GSTIN: {gstin}\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\tContact: +91 {contact}')
        header_run.font.name = 'Calibri'
        header_run.font.size = Pt(12)
        header_run.bold = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Main heading
        main_heading = doc.add_paragraph()
        main_heading_run = main_heading.add_run(business_name)
        main_heading_run.font.name = 'Adobe Garamond Pro Bold'
        main_heading_run.font.size = Pt(40)
        main_heading_run.bold = True
        main_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        main_heading.paragraph_format.space_after = Pt(9)
        
        # Document type and date in separate paragraphs
        doc_type_para = doc.add_paragraph()
        doc_type_run = doc_type_para.add_run(doc_type)
        doc_type_run.font.name = 'Calibri'
        doc_type_run.font.size = Pt(24)
        doc_type_run.underline = True
        doc_type_run.bold = True
        doc_type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'Date: {formatted_date}')
        date_run.font.name = 'Calibri Light (Headings)'
        date_run.font.size = Pt(16)
        date_run.bold = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Customer info
        customer_para = doc.add_paragraph()
        customer_run1 = customer_para.add_run("To,")
        customer_run1.font.size = Pt(16)
        customer_run1.font.name = "Calibri"
        customer_run1.bold = True
        customer_run2 = customer_para.add_run(f"\n{customer_name}")
        customer_run2.font.size = Pt(16)
        customer_run2.font.name = "Calibri"
        customer_run2.bold = True
        
        # Create table for items with columns
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Item'
        header_cells[1].text = 'Qty'
        header_cells[2].text = 'Rate'
        header_cells[3].text = 'Amount'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(12)
        
        # Add items to table
        for i, item in enumerate(st.session_state["items"], 1):
            row_cells = table.add_row().cells

            spec_text = f": {item.get('specification', '')}" if item.get('specification') else ""
            row_cells[0].text = f"{i}. {item['name']}{spec_text}"

            mode = item.get("mode", "Quantity + Rate")
            qty_text = ""
            rate_text = ""
            amt_text = f"₹{item.get('total', 0):,.2f}"

            # Quantity + Rate
            if mode == "Quantity + Rate":
                qty_text = f"{item.get('quantity', 0)} {item.get('unit', '')}"
                rate_text = f"₹{item.get('rate', 0):,.2f}/{item.get('unit', '')}"

            # Quantity + Amount
            elif mode == "Quantity + Amount":
                qty_text = f"{item.get('quantity', 0)} {item.get('unit', '')}"
                rate_text = "—"

            # Rate + Amount
            elif mode == "Rate + Amount":
                qty_text = "—"
                rate_text = f"₹{item.get('rate', 0):,.2f}"

            # Amount Only
            elif mode == "Amount Only":
                qty_text = "—"
                rate_text = "—"

            row_cells[1].text = qty_text
            row_cells[2].text = rate_text
            row_cells[3].text = amt_text

        # Subtotal and GST rows
        subtotal = sum(item['total'] for item in st.session_state["items"]) if st.session_state["items"] else 0.0

        # Compute GST exactly like in the UI
        gst_percent = 0.0 if gst_choice == "None" else float(gst_choice.replace("%", ""))
        if gst_percent > 0:
            if gst_mode.startswith("Exclusive"):
                gst_amount = subtotal * gst_percent / 100.0
                amount_chargeable = subtotal + gst_amount
            else:
                gst_amount = subtotal * gst_percent / (100.0 + gst_percent)
                amount_chargeable = subtotal
            cgst = gst_amount / 2.0
            sgst = gst_amount / 2.0
        else:
            gst_amount = cgst = sgst = 0.0
            amount_chargeable = subtotal

        # GST rows (insert before final totals)
        if gst_percent > 0:
            # CGST
            cgst_row = table.add_row().cells
            cgst_row[0].text = ""
            cgst_row[1].text = ""
            cgst_row[2].text = f"CGST @ {gst_percent/2:.2f}%"
            cgst_row[3].text = f"₹{cgst:,.2f}"
            # SGST
            sgst_row = table.add_row().cells
            sgst_row[0].text = ""
            sgst_row[1].text = ""
            sgst_row[2].text = f"SGST @ {gst_percent/2:.2f}%"
            sgst_row[3].text = f"₹{sgst:,.2f}"

        # Total Amount row
        total_row = table.add_row().cells
        total_row[0].text = ""
        total_row[1].text = ""
        total_row[2].text = "Total Amount"
        total_row[3].text = f"₹{amount_chargeable:,.2f}"
        for cell in total_row[2:]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = 'Calibri'
                    r.font.size = Pt(11)

        # Advance and Balance rows
        if advance > 0:
            adv_row = table.add_row().cells
            adv_row[0].text = ""
            adv_row[1].text = ""
            adv_row[2].text = "Advance"
            adv_row[3].text = f"₹{advance:,.2f}"

            bal_row = table.add_row().cells
            bal_row[0].text = ""
            bal_row[1].text = ""
            bal_row[2].text = "Balance"
            bal_row[3].text = f"₹{max(0.0, amount_chargeable - advance):,.2f}"

            # Make both rows bold
            for row in [adv_row, bal_row]:
                for cell in row[2:]:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = 'Calibri'

        # Align all text in the last column (amounts) to the right
        for row in table.rows:
            last_cell = row.cells[-1]
            for paragraph in last_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Material used
        # -----------------------------
        # WRITE MULTIPLE MATERIAL SECTIONS
        # -----------------------------
        if "materials" in st.session_state and st.session_state["materials"]:
            # add a blank paragraph separator
            doc.add_paragraph()

            for mat in st.session_state["materials"]:
                # Header for this material block (e.g., "Material Used: SS ...")
                material_full = f"{mat.get('type','')} {mat.get('grade','')}".strip()
                material_header = doc.add_paragraph()
                material_header_run = material_header.add_run(f"Material Used: {material_full}")
                material_header_run.bold = True
                material_header_run.font.name = 'Calibri'
                material_header_run.font.size = Pt(14)

                # Proper bullet list using Word's built-in style
                for detail in mat.get("details", []):
                    if detail.strip():
                        material_para = doc.add_paragraph(detail.strip(), style='List Bullet')
                        for run in material_para.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(12)


        # -----------------------------
        # 📜 TERMS & CONDITIONS
        # -----------------------------
        if selected_terms:
            doc.add_paragraph()
            terms_header = doc.add_paragraph()
            terms_header_run = terms_header.add_run('Terms & Conditions:')
            terms_header_run.bold = True
            terms_header_run.font.name = 'Calibri'
            terms_header_run.font.size = Pt(12)

            # Proper bullet list using Word's native style
            for term in selected_terms:
                if term.strip():
                    term_para = doc.add_paragraph(term.strip(), style='List Bullet')
                    for run in term_para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
        
        return doc

    if st.button("📥 Download Word Document", type="secondary"):
        if customer_name and st.session_state["items"]:
            doc = create_word_document()
            
            # Save to bytes
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Download button
            st.download_button(
                label="📄 Download Document",
                data=doc_buffer.getvalue(),
                file_name=f"{doc_type}_{customer_name}_{formatted_date.replace('/', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Please enter customer name and add at least one item")

# Clear all button
if st.button("🗑️ Clear All Data", type="secondary"):
    st.session_state["items"] = []
    st.rerun()

# Footer
st.markdown("---")
st.markdown("*Quotation Generator v1.0*")